import os
import logging
import docx
import pandas as pd
import asyncio

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

try:
    import tools.sbir as sbir
    import tools.phase2 as phase2
    import tools.grading as grading
except ImportError as e:
    logging.error(f"[SBIR PIPELINE] Could not import a phase script: {e}")

    async def run_and_read_sbir_pipeline():
        logging.error("[SBIR PIPELINE] Pipeline scripts missing. Returning empty result.")
        return []

    def fetch_sbir_partnership_opportunities():
        return []
else:
    async def run_and_read_sbir_pipeline():
        logging.info("[SBIR PIPELINE] Running 3-phase SBIR pipeline...")
        try:
            sbir.run_phase_1()
            await phase2.run_phase_2()
            await grading.run_grading_process()
            logging.info("[SBIR PIPELINE] Word document generation complete.")
        except Exception as e_pipeline:
            logging.error(f"[SBIR PIPELINE] Error during pipeline execution: {e_pipeline}", exc_info=True)
            return []

        logging.info("[SBIR PIPELINE] Reading generated Word dossiers...")
        dossier_folder = 'company_dossiers'
        if not os.path.isdir(dossier_folder):
            logging.error(f"[SBIR PIPELINE] Dossier folder '{dossier_folder}' not found.")
            return []

        partner_list = []
        for filename in os.listdir(dossier_folder):
            if filename.endswith('.docx'):
                try:
                    doc_path = os.path.join(dossier_folder, filename)
                    doc = docx.Document(doc_path)
                    doc_text = "\n".join([p.text for p in doc.paragraphs])
                    company_name = "Unknown Company"
                    project_title = filename.replace('.docx', '')
                    for para in doc.paragraphs:
                        line = para.text.strip()
                        if line.lower().startswith("company name:"):
                            company_name = line.split(":", 1)[1].strip()
                        elif line.lower().startswith("project title:"):
                            project_title = line.split(":", 1)[1].strip()
                    partner_list.append({
                        "company_name": company_name,
                        "project_title": project_title,
                        "full_text": doc_text
                    })
                except Exception as e_doc:
                    logging.warning(f"[SBIR PIPELINE] Could not parse '{filename}': {e_doc}")

        logging.info(f"[SBIR PIPELINE] Parsed {len(partner_list)} dossiers.")
        return partner_list

    def fetch_sbir_partnership_opportunities():
        return asyncio.run(run_and_read_sbir_pipeline())
