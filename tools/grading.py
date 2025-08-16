# grading.py
import os
import pandas as pd
import logging
import asyncio
import json
import re
import random
from datetime import datetime
from typing import Dict, Optional, Any, List, Tuple

# Document and file handling imports
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import httpx
from openai import AsyncAzureOpenAI, APITimeoutError, RateLimitError

# ====== CONFIGURATION & CONSTANTS ======
AZURE_OPENAI_ENDPOINT = (os.getenv("AZURE_OPENAI_ENDPOINT") or "").strip()
AZURE_OPENAI_KEY = (os.getenv("AZURE_OPENAI_KEY") or "").strip()
AZURE_OPENAI_DEPLOYMENT = (os.getenv("AZURE_OPENAI_DEPLOYMENT") or "").strip()

DOSSIER_FOLDER = "company_dossiers"
OUTPUT_FILENAME = "sbir_top_30_ranked.docx"
COMPANY_DATA_SOURCE_FILE = "Discovered Companies.xlsx"
LOG_FILENAME = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sbir_grading_log.txt")

MAX_RANKING_LIMIT = 30
GRADING_WEIGHTS: Dict[str, float] = {
    "Technology_Strength": 0.30,
    "Market_Traction": 0.30,
    "Team_Experience": 0.20,
    "DoD_Alignment": 0.20,
}

# ====== HTTP/CLIENT POOLING ======
MAX_CONNECTIONS = 4
CONCURRENCY = 2                 # keep modest to reduce 429s
CONNECT_TIMEOUT = 30.0
READ_TIMEOUT = 120.0
POOL_TIMEOUT = 60.0
USE_HTTP2 = False               # flip to True only if your Azure endpoint supports it
GRADING_BATCH_SIZE = 4          # used by the batching loop below

_HTTP_LIMITS = httpx.Limits(
    max_connections=MAX_CONNECTIONS,
    max_keepalive_connections=MAX_CONNECTIONS,
)
_HTTP_TIMEOUT = httpx.Timeout(
    connect=CONNECT_TIMEOUT,
    read=READ_TIMEOUT,
    write=60.0,
    pool=POOL_TIMEOUT,
)
_httpx_async_client = httpx.AsyncClient(
    limits=_HTTP_LIMITS,
    timeout=_HTTP_TIMEOUT,
    http2=USE_HTTP2,
)

_client: Optional[AsyncAzureOpenAI] = None
_SEMAPHORE = asyncio.Semaphore(CONCURRENCY)

def _extract_json_lenient(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if not m:
        return None
    s = m.group(0)
    # sanitize common artifacts
    s = re.sub(r"\{\s*\.\.\.\s*\}", "{}", s)
    s = re.sub(r"\[\s*\.\.\.\s*\]", "[]", s)
    s = s.replace("[calculated value]", "0")
    s = re.sub(r"\bNaN\b", "0", s)
    s = re.sub(r"(?m)^\s*//.*$", "", s)
    try:
        return json.loads(s)
    except Exception:
        return None
    
def _retry_after_seconds_from_429_message(msg: str) -> float:
    # Azure often says: "Please retry after X seconds."
    m = re.search(r"retry after (\d+) second", msg, re.IGNORECASE)
    if m:
        try:
            return float(m.group(1))
        except Exception:
            pass
    # default backoff if not present
    return 2.0

# ====== AI GRADING FUNCTION (RETRY & THROTTLE) ======
async def get_ai_grading(dossier_text: str) -> Dict[str, Any]:
    """Asynchronously grades a company dossier with backoff and strict JSON parsing."""
    assert AZURE_OPENAI_DEPLOYMENT, "AZURE_OPENAI_DEPLOYMENT is required"

    system_prompt = (
        "You are a lead analyst on a venture capital investment committee, specializing in DoD technology. "
        "Return ONLY a valid JSON object and nothing else."
    )
    user_prompt = f"""
Analyze the following company dossier and provide a quantitative and qualitative assessment.
Return your response ONLY as a valid JSON object with the following structure:
{{
  "Technology_Strength": {{"score": 0.0, "justification": "..." }},
  "Market_Traction": {{"score": 0.0, "justification": "..." }},
  "Team_Experience": {{"score": 0.0, "justification": "..." }},
  "DoD_Alignment": {{"score": 0.0, "justification": "..." }}
}}

Dossier Text:
---
{dossier_text}
""".strip()

    async with _SEMAPHORE:
        backoff = 0.8
        max_attempts = 6
        for attempt in range(max_attempts):
            try:
                assert _client is not None
                resp = await _client.chat.completions.create(
                    model=AZURE_OPENAI_DEPLOYMENT,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    temperature=0.1,
                    max_tokens=700,
                    timeout=READ_TIMEOUT,
                    response_format={"type": "json_object"},
                )
                raw_content = (resp.choices[0].message.content or "").strip()
                if not raw_content:
                    raise RuntimeError("AI returned an empty response.")
                
                raw_preview = raw_content[:400].replace("\n", " ")
                logging.debug("AI raw preview: %s%s", raw_preview, "..." if len(raw_content) > 400 else "")

                # strict parse first
                try:
                    m = re.search(r"\{.*\}", raw_content, re.DOTALL)
                    if m:
                        return json.loads(m.group(0))
                except json.JSONDecodeError:
                    pass

                parsed = _extract_json_lenient(raw_content)
                if parsed is not None:
                    return parsed

                raise ValueError("AI returned a non-JSON or malformed JSON response.")

            except RateLimitError as e:
                # Respect server hint and continue backoff curve
                msg = str(e)
                sleep_for = max(_retry_after_seconds_from_429_message(msg), 1.0)
                logging.warning("429 rate limit on grading (attempt %d/%d). Sleeping %.1fs. %s",
                                attempt + 1, max_attempts, sleep_for, e)
                if attempt == max_attempts - 1:
                    logging.error("Final 429 on grading: %s", e, exc_info=True)
                    return {"error": "Rate limited."}
                await asyncio.sleep(sleep_for)
                backoff *= 1.6

            except (APITimeoutError, httpx.ReadTimeout, httpx.ConnectTimeout, httpx.PoolTimeout) as e:
                logging.warning("Transient timeout on grading (attempt %d/%d): %s",
                                attempt + 1, max_attempts, e)
                if attempt == max_attempts - 1:
                    logging.error("Final timeout on grading: %s", e, exc_info=True)
                    return {"error": "Request timed out."}
                sleep_for = backoff + random.uniform(0, 0.5)
                await asyncio.sleep(sleep_for)
                backoff *= 1.8

            except httpx.HTTPError as e:
                logging.warning("HTTP error on grading (attempt %d/%d): %s",
                                attempt + 1, max_attempts, e)
                if attempt == max_attempts - 1:
                    logging.error("HTTP error on grading: %s", e, exc_info=True)
                    return {"error": str(e)}
                sleep_for = backoff + random.uniform(0, 0.5)
                await asyncio.sleep(sleep_for)
                backoff *= 1.8

            except Exception as e:
                # Log the first ~400 chars of content to diagnose non-JSON (but don't crash the batch)
                logging.error("AI grading failed (non-retryable): %s", e, exc_info=True)
                return {"error": "AI returned a non-JSON response."}

    # Safety net so Pylance sees a return on all paths
    return {"error": "Unknown failure after retries."}

# ====== HELPER FUNCTIONS ======
def load_company_urls(source_file: str) -> Dict[str, str]:
    try:
        if not os.path.exists(source_file):
            logging.warning("Company data source file not found at: %s", source_file)
            return {}
        df = pd.read_excel(source_file)
        url_col = next((col for col in ["company_url", "firm_company_url"] if col in df.columns), None)
        if "firm" in df.columns and url_col:
            df[url_col] = df[url_col].fillna("N/A")
            return df.drop_duplicates(subset="firm").set_index("firm")[url_col].to_dict()
        return {}
    except Exception as e:
        logging.error("Error loading company URLs: %s", e)
        return {}

def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def save_graded_report(df: pd.DataFrame, filename: str) -> None:
    if df.empty:
        logging.warning("No graded data to save.")
        return
    df_sorted = df.sort_values(by="Smart_Bet_Score", ascending=False).head(MAX_RANKING_LIMIT)
    doc = Document()
    doc.add_heading("SBIR Company Grading & Ranking Report", level=0)
    p_date = doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for rank, (idx, row) in enumerate(df_sorted.iterrows(), 1):
        doc.add_heading(f"Rank #{rank}: {row['Company']}", level=1)
        url = row.get("Company_URL", "")
        if isinstance(url, str) and url.startswith("http"):
            p_url = doc.add_paragraph()
            add_hyperlink(p_url, url, url)
        p_score = doc.add_paragraph()
        r = p_score.add_run("Overall Smart Bet Score: ")
        r.bold = True
        p_score.add_run(f"{row['Smart_Bet_Score']:.2f}").font.size = Pt(14)
        doc.add_heading("Detailed Assessment", level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = "Category", "Score", "Justification"
        for cat in ["Technology_Strength", "Market_Traction", "Team_Experience", "DoD_Alignment"]:
            row_cells = table.add_row().cells
            row_cells[0].text = cat.replace("_", " ")
            row_cells[1].text = str(row.get(f"{cat}_Score", "N/A"))
            row_cells[2].text = str(row.get(f"{cat}_Justification", "N/A"))
        doc.add_paragraph()
    doc.save(filename)
    logging.info("Saved graded report to %s", filename)

def _to_float(val: Any) -> float:
    try:
        return float(val)
    except Exception:
        return 0.0

def safe_get_score(score_dict: Optional[Dict[str, Any]]) -> float:
    if isinstance(score_dict, dict) and "score" in score_dict:
        return _to_float(score_dict["score"])
    return 0.0

def safe_get_justification(score_dict: Optional[Dict[str, Any]]) -> str:
    if isinstance(score_dict, dict) and "justification" in score_dict:
        return str(score_dict["justification"])
    return "N/A"

# ====== MAIN ORCHESTRATOR ======
async def run_grading_process() -> None:
    if not os.path.isdir(DOSSIER_FOLDER):
        logging.error("Dossier folder does not exist: %s", DOSSIER_FOLDER)
        return

    # Validate Azure credentials early
    if not (AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_KEY and AZURE_OPENAI_DEPLOYMENT):
        logging.error("Azure OpenAI credentials not fully set. Halting grading process.")
        return

    # Init client once (reuse pooled httpx client)
    global _client
    _client = AsyncAzureOpenAI(
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        api_key=AZURE_OPENAI_KEY,
        api_version="2024-02-01",
        http_client=_httpx_async_client,
    )

    company_urls = load_company_urls(COMPANY_DATA_SOURCE_FILE)
    all_graded_results: List[Dict[str, Any]] = []

    dossier_files = glob.glob(os.path.join(DOSSIER_FOLDER, "*.docx"))
    logging.info("Found %d dossiers to process.", len(dossier_files))

    # === Build inputs first ===
    entries: List[Tuple[str, str]] = []  # (company_name, full_text)
    company_data_map: Dict[str, Dict[str, str]] = {}

    for file_path in dossier_files:
        try:
            doc = Document(file_path)
            if not doc.paragraphs:
                continue
            company_name = doc.paragraphs[0].text.replace("Dossier:", "").strip()
            full_text = "\n".join(p.text for p in doc.paragraphs)
            entries.append((company_name, full_text))
            company_data_map[company_name] = {"url": company_urls.get(company_name, "N/A")}
        except Exception as e:
            logging.error("Error reading dossier %s: %s", file_path, e)

    # === Process in small batches to avoid 429s ===
    batch_size = GRADING_BATCH_SIZE
    names_order: List[str] = []
    ai_grades: List[Any] = []

    for i in range(0, len(entries), batch_size):
        batch = entries[i:i + batch_size]
        names_order.extend([nm for nm, _ in batch])

        batch_tasks = [
            asyncio.create_task(get_ai_grading(txt))
            for (_, txt) in batch
        ]
        batch_results = await asyncio.gather(*batch_tasks, return_exceptions=True)
        ai_grades.extend(batch_results)

    for company_name, grade in zip(names_order, ai_grades):
        data = company_data_map.get(company_name, {"url": "N/A"})

        if isinstance(grade, Exception):
            logging.error("Failed grading for %s: %s", company_name, grade)
            continue
        if not isinstance(grade, dict) or "error" in grade:
            logging.error("Failed grading for %s: %s", company_name, getattr(grade, "error", grade))
            continue

        tech_score = safe_get_score(grade.get("Technology_Strength"))
        market_score = safe_get_score(grade.get("Market_Traction"))
        team_score = safe_get_score(grade.get("Team_Experience"))
        dod_score = safe_get_score(grade.get("DoD_Alignment"))

        smart_bet = (
            (tech_score * GRADING_WEIGHTS["Technology_Strength"])
            + (market_score * GRADING_WEIGHTS["Market_Traction"])
            + (team_score * GRADING_WEIGHTS["Team_Experience"])
            + (dod_score * GRADING_WEIGHTS["DoD_Alignment"])
        )

        all_graded_results.append({
            "Company": company_name,
            "Company_URL": data["url"],
            "Smart_Bet_Score": round(smart_bet, 2),
            "Technology_Strength_Score": tech_score,
            "Technology_Strength_Justification": safe_get_justification(grade.get("Technology_Strength")),
            "Market_Traction_Score": market_score,
            "Market_Traction_Justification": safe_get_justification(grade.get("Market_Traction")),
            "Team_Experience_Score": team_score,
            "Team_Experience_Justification": safe_get_justification(grade.get("Team_Experience")),
            "DoD_Alignment_Score": dod_score,
            "DoD_Alignment_Justification": safe_get_justification(grade.get("DoD_Alignment")),
        })
        logging.info("Graded %s with score %.2f", company_name, smart_bet)

    if all_graded_results:
        df = pd.DataFrame(all_graded_results)
        save_graded_report(df, OUTPUT_FILENAME)

# Clean shutdown for pooled clients
import atexit
@atexit.register
def _close_clients():
    try:
        if not _httpx_async_client.is_closed:
            # Close asynchronously even if event loop isn’t running
            try:
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    loop.create_task(_httpx_async_client.aclose())
                else:
                    loop.run_until_complete(_httpx_async_client.aclose())
            except RuntimeError:
                # No loop; create a temporary one
                asyncio.run(_httpx_async_client.aclose())
    except Exception:
        pass

if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[logging.FileHandler(LOG_FILENAME, mode="w"), logging.StreamHandler()],
    )
    asyncio.run(run_grading_process())
