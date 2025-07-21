import asyncio
import os
import pandas as pd
import time
import logging
from datetime import datetime
import json
import random
from docx import Document
import re
from dotenv import load_dotenv

from azure.ai.inference.aio import ChatCompletionsClient
from azure.ai.inference.models import SystemMessage, UserMessage
from azure.core.credentials import AzureKeyCredential

# --- Load ENV ---
load_dotenv()

AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")
AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")

if not (AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_KEY and AZURE_OPENAI_DEPLOYMENT):
    raise EnvironmentError("Azure OpenAI ENV variables missing.")

INPUT_FILENAME = "Discovered Companies.xlsx"
OUTPUT_FOLDER = "company_dossiers"
LOG_FILENAME = "sbir_research_log.txt"

def clean_ai_response(text):
    patterns_to_remove = [
        r"^\s*Okay,.*?dossier\s*\n",
        r"^\s*Here is the dossier:?\s*\n",
        r"^\s*Okay, initiating.*?Complete\.\.\.\s*\n",
        r"^\s*Okay, here's.*?dossier\s*\n",
        r"^\s*Okay, here is my analysis of.*?as requested\.\s*\n",
        r"^\s*Okay, here's the venture capital analyst dossier.*?\n"
    ]
    cleaned_text = text
    for pattern in patterns_to_remove:
        cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE | re.DOTALL)
    return cleaned_text.strip()

async def get_ai_research_summary(company_name):
    if not AZURE_OPENAI_KEY or not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_DEPLOYMENT:
        return "AI summary requires valid Azure OpenAI environment variables."

    client = ChatCompletionsClient(
        endpoint=AZURE_OPENAI_ENDPOINT,
        credential=AzureKeyCredential(AZURE_OPENAI_KEY),
    )

    prompt = f"""
    Act as a senior venture capital analyst specializing in defense and aerospace. Research the US-based company '{company_name}'.

    **Company Overview:** One paragraph overview of business, mission, and value proposition.

    **Technology Focus:** 2-3 bullet points on technology/products.

    **Recent Developments & Traction:** 2-4 bullet points on recent news, partnerships, funding.

    **Leadership & Team:** Key leaders and notable past experience.

    **Competitive Landscape:** 1-2 competitors and differentiation.

    **Sources:** Top 3-5 URLs.
    """

    try:
        result = await client.complete(
            deployment_name=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                SystemMessage(content="You are a venture analyst assistant."),
                UserMessage(content=prompt)
            ],
            temperature=0.4,
            max_tokens=2000
        )
        ai_response = result.choices[0].message.content
        return clean_ai_response(ai_response)

    except Exception as e:
        logging.error(f"Azure OpenAI API Error for {company_name}: {e}")
        return "AI summary failed due to API error."

def create_company_dossier(company_name, award_data, ai_summary):
    doc = Document()
    doc.add_heading(f'Dossier: {company_name}', level=1)

    doc.add_heading('SBIR Award Details', level=2)
    p = doc.add_paragraph()
    p.add_run('Award Title: ').bold = True
    p.add_run(str(award_data.get('award_title', 'N/A')))

    p = doc.add_paragraph()
    p.add_run('Amount: ').bold = True
    p.add_run(f"${float(award_data.get('award_amount', 0)):,.2f}")

    p = doc.add_paragraph()
    p.add_run('Award Date: ').bold = True
    award_date = award_data.get('proposal_award_date', 'N/A')
    if isinstance(award_date, datetime):
        p.add_run(award_date.strftime('%Y-%m-%d'))
    else:
        p.add_run(str(award_date))

    p = doc.add_paragraph()
    p.add_run('Branch: ').bold = True
    p.add_run(str(award_data.get('branch', 'N/A')))

    doc.add_heading('AI-Generated Intelligence Summary', level=2)
    for line in ai_summary.split('\n'):
        line = line.strip()
        if line.startswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '').strip()).bold = True

        elif line.startswith('*'):
            doc.add_paragraph(line.lstrip('* ').strip(), style='List Bullet')
        elif line:
            doc.add_paragraph(line)

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    safe_filename = "".join(c for c in company_name if c.isalnum() or c == ' ').rstrip()
    filepath = os.path.join(OUTPUT_FOLDER, f"{safe_filename}.docx")
    doc.save(filepath)
    logging.info(f"Saved dossier for {company_name} at {filepath}")

async def run_research_and_generate_dossiers(input_file):
    logging.info(f"Starting AI research from {input_file}")
    try:
        df = pd.read_excel(input_file)
    except FileNotFoundError:
        logging.error(f"File not found: {input_file}")
        return

    companies = df.drop_duplicates(subset='firm')
    logging.info(f"Found {len(companies)} unique companies")

    for _, row in companies.iterrows():
        company_name = row['firm']
        logging.info(f"Processing: {company_name}")
        summary = await get_ai_research_summary(company_name)
        create_company_dossier(company_name, row, summary)
        await asyncio.sleep(2)

async def run_phase_2():
    logging.info("--- Phase 2 Started ---")
    await run_research_and_generate_dossiers(INPUT_FILENAME)
    logging.info("--- Phase 2 Complete ---")

if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[logging.FileHandler(LOG_FILENAME), logging.StreamHandler()]
    )
    asyncio.run(run_phase_2())
